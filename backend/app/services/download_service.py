import json
from pathlib import Path

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
)
from reportlab.lib.styles import (
    getSampleStyleSheet,
)

from docx import Document

from openpyxl import Workbook


def generate_download_file(
    content: str,
    file_type: str,
    output_path: str,
):
    """
    Generate downloadable file from text.

    Supported:
    - pdf
    - docx
    - md
    - json
    - xlsx
    """

    output_path = Path(output_path)

    if file_type == "pdf":

        doc = SimpleDocTemplate(
            str(output_path)
        )

        styles = (
            getSampleStyleSheet()
        )

        story = [
            Paragraph(
                content.replace(
                    "\n",
                    "<br/>"
                ),
                styles["BodyText"],
            )
        ]

        doc.build(story)

    elif file_type == "docx":

        document = Document()

        document.add_heading(
            "Generated Answer",
            level=1,
        )

        document.add_paragraph(
            content
        )

        document.save(
            str(output_path)
        )

    elif file_type == "md":

        with open(
            output_path,
            "w",
            encoding="utf-8",
        ) as f:

            f.write(content)

    elif file_type == "json":

        try:

            parsed = json.loads(
                content
            )

        except Exception:

            parsed = {
                "answer": content
            }

        with open(
            output_path,
            "w",
            encoding="utf-8",
        ) as f:

            json.dump(
                parsed,
                f,
                indent=2,
                ensure_ascii=False,
            )

    elif file_type == "xlsx":

        wb = Workbook()

        ws = wb.active

        ws.title = "Answer"

        lines = content.split(
            "\n"
        )

        ws.cell(
            row=1,
            column=1,
            value="Generated Answer",
        )

        for idx, line in enumerate(
            lines,
            start=2,
        ):

            ws.cell(
                row=idx,
                column=1,
                value=line,
            )

        wb.save(
            str(output_path)
        )

    else:

        raise ValueError(
            f"Unsupported file type: {file_type}"
        )